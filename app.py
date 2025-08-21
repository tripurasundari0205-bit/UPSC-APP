# app.py
# ---------------------------------------------
# UPSC PYQ → Structured Solution Book Generator
# Hosted Streamlit app with DOCX/PDF export
# ---------------------------------------------

import os
import re
import io
import json
import uuid
from dataclasses import dataclass
from typing import List, Optional, Dict

import streamlit as st
import pdfplumber

# Export libs
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet

# LLM (OpenAI SDK)
from openai import OpenAI
from groq import Groq
# ---------------------------
# App Config
# ---------------------------
st.set_page_config(page_title="UPSC Solution Generator", page_icon="📘", layout="wide")
st.title("📘 UPSC PYQ → Structured Solution Generator")
st.caption("Upload a PYQ PDF and generate exam-ready solutions with citations, then export as DOCX/PDF.")

# ---------------------------
# Constants & Prompts
# ---------------------------
ALLOWED_DOMAINS = [
    ".gov.in", "pib.gov.in", "prsindia.org", "indiacode.nic.in", "rbi.org.in",
    "sebi.gov.in", "moef.gov.in", "niti.gov.in", "censusindia.gov.in",
    "un.org", "worldbank.org", "imf.org", "ipcc.ch", "undp.org", "unesco.org",
    "who.int", "fao.org", "oecd.org", "ncert.nic.in", "cag.gov.in"
]

SYSTEM_PROMPT = """You are an expert UPSC faculty member. Produce solutions STRICTLY in this structure:

Question: <verbatim question with options>

1. Correct Answer – State the correct option (A/B/C/D or ‘Not applicable’) and a one-sentence reason.

2. Concept Behind the Question – Explain the background, core concept, related theory, and why it matters for UPSC.

3. Explanation of Each Option – For every option, say “Correct” or “Incorrect” with 2–4 lines of reasoning and the concept behind it.

4. Relevance/Current Affairs Linkage – Link to present-day relevance, schemes, reports, or applications. If none, say “No direct current linkage.”

5. Quick Summary for Revision – 4–6 crisp bullet points.

Rules:
- Be accurate and concise. Avoid speculation.
- Cite only reliable sources. List sources at the end (bulleted).
- If unsure about correctness, say so and suggest verification.
- Keep tone exam-oriented; avoid fluff.
"""

USER_TEMPLATE = """PDF_Source_Metadata: {meta}

Question_Block:
{question_block}

Detected_Type: {qtype}

Topic_Tags: {tags}

Retrieved_Context:
{retrieved}
"""

# ---------------------------
# Data models
# ---------------------------
@dataclass
class Question:
    qid: str
    text: str
    options: List[str]
    qtype: str   # MCQ | Mains | Assertion-Reason | etc.
    year: Optional[str] = None
    paper: Optional[str] = None
    topic_tags: Optional[List[str]] = None
    page: Optional[int] = None

@dataclass
class Solution:
    qid: str
    content_markdown: str
    citations: List[str]

# ---------------------------
# PDF parsing
# ---------------------------
Q_START = re.compile(r"^\s*(\d{1,3})\s*[).:-]\s+", re.MULTILINE)
OPT_LINE = re.compile(r"^\s*[A-D]\s*[\)\].-]\s+", re.MULTILINE)

def extract_text_from_pdf(file) -> str:
    text_parts = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            # remove common headers/footers heuristically
            lines = [ln for ln in t.splitlines() if len(ln.strip()) > 0]
            text_parts.append("\n".join(lines))
    return "\n\n".join(text_parts)

def split_questions(raw: str) -> List[str]:
    # Split at question numbers while keeping the headers
    indices = [m.start() for m in Q_START.finditer(raw)]
    blocks = []
    if not indices:
        return blocks
    indices.append(len(raw))
    for i in range(len(indices)-1):
        blocks.append(raw[indices[i]:indices[i+1]].strip())
    return blocks

def parse_mcq_block(block: str) -> Question:
    # Remove leading number "12) " etc.
    block_clean = re.sub(r"^\s*\d{1,3}\s*[).:-]\s*", "", block).strip()
    # Find options
    opt_matches = list(OPT_LINE.finditer(block_clean))
    options = []
    if opt_matches:
        # Split by option markers
        parts = OPT_LINE.split(block_clean)
        stem = parts[0].strip()  # question stem
        # Recreate labeled options in order A-D
        labels = [re.findall(r"[A-D]", block_clean[m.start():m.end()])[0] for m in opt_matches]
        for label_letter, p in zip(labels, parts[1:]):
            option_text = p.strip()
            options.append(f"{label_letter}. {option_text}")
        qtype = "MCQ"
    else:
        stem = block_clean
        qtype = "Mains"
    return Question(
        qid=str(uuid.uuid4())[:8],
        text=stem,
        options=options,
        qtype=qtype
    )

# ---------------------------
# Topic tagging (simple heuristic; replace with LLM if desired)
# ---------------------------
KEYWORDS_TO_TAGS = {
    "parliament": "Polity", "governor": "Polity", "fundamental right": "Polity",
    "inflation": "Economy", "gdp": "Economy", "rbi": "Economy",
    "mangrove": "Environment", "biosphere": "Environment", "unfccc": "Environment",
    "dna": "Science & Tech", "quantum": "Science & Tech",
    "harappan": "History", "buddh": "History",
    "map": "Geography", "monsoon": "Geography",
}

def tag_topics(text: str) -> List[str]:
    lower = text.lower()
    tags = set()
    for k, v in KEYWORDS_TO_TAGS.items():
        if k in lower:
            tags.add(v)
    return list(tags) or ["General Studies"]

# ---------------------------
# Retrieval stub (optional)
# ---------------------------
def retrieve_snippets(question: Question) -> List[str]:
    """
    Placeholder for retrieval. Options:
      - Local vector store over NCERT/standard notes
      - Web search API restricted to ALLOWED_DOMAINS
    Return bullets like: "- Fact ... [PIB, 2023]"
    """
    return []  # start empty; upgrade later

# ---------------------------
# LLM call (OpenAI, reads key from Streamlit secrets)
# ---------------------------
import streamlit as st
from openai import OpenAI
from groq import Groq

# --- Updated call_llm with fallback ---
def call_llm(system_prompt: str, user_prompt: str) -> str:
    try:
        # Try OpenAI first
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
        )
        return resp.choices[0].message.content

    except Exception as e:
        st.warning(f"⚠️ OpenAI failed ({e}), switching to Groq fallback...")

        try:
            groq_client = Groq(api_key=st.secrets["GROQ_API_KEY"])
            resp = groq_client.chat.completions.create(
                model="hog2-wizard-15b",  # strong model, 32k context
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2,
            )
            return resp.choices[0].message.content

        except Exception as e2:
            st.error(f"❌ Both OpenAI and Groq failed: {e2}")
            return "Error: Unable to generate response at this time."


# ---------------------------
# Validation
# ---------------------------
def validate_markdown(md: str, q: Question) -> List[str]:
    errs = []

    # Section checks
    required_sections = [
        "Question:",
        "1. Correct Answer",
        "2. Concept Behind the Question",
        "3. Explanation of Each Option",
        "4. Relevance/Current Affairs Linkage",
        "5. Quick Summary for Revision",
    ]
    for head in required_sections:
        if head.split(":")[0] not in md:
            errs.append(f"Missing section: {head}")

    # MCQ: ensure one correct option stated (A-D)
    if q.qtype == "MCQ":
        m = re.search(r"1\.\s*Correct Answer\s*[\-–]\s*.*\b([ABCD])\b", md, flags=re.IGNORECASE)
        if not m:
            errs.append("Correct Answer section missing A/B/C/D label.")
        # Ensure each present option is at least discussed
        for opt in q.options:
            label = opt.split(".")[0].strip()
            # Look for mentions like "Option A" or "A)"/"A."
            if not re.search(rf"(Option\s*{label}\b|\b{label}\)|\b{label}\.)", md, flags=re.IGNORECASE):
                # soft warning only; different phrasing possible
                pass

    # Citations: expect "Sources" or "References"
    if ("Sources" not in md) and ("References" not in md):
        errs.append("Missing sources/citations section.")

    return errs

# ---------------------------
# Orchestration
# ---------------------------
def generate_solution(q: Question) -> Solution:
    retrieved = retrieve_snippets(q)
    user_prompt = USER_TEMPLATE.format(
        meta=json.dumps({"year": q.year, "paper": q.paper, "page": q.page}),
        question_block=(q.text + ("\n\n" + "\n".join(q.options) if q.options else "")),
        qtype=q.qtype,
        tags=", ".join(q.topic_tags or []),
        retrieved="\n".join("- " + s for s in retrieved) if retrieved else "None"
    )
    md = call_llm(SYSTEM_PROMPT, user_prompt)
    errs = validate_markdown(md, q)
    if errs:
        md += "\n\n---\nValidator notes (for editor):\n" + "\n".join(f"- {e}" for e in errs)
    # naive citation scrape from any URLs
    cites = re.findall(r"\((https?://[^\s)]+)\)", md)
    return Solution(qid=q.qid, content_markdown=md, citations=cites)

# ---------------------------
# Export helpers
# ---------------------------
def export_docx(solutions: List[Solution], title: str = "UPSC Solution Book") -> bytes:
    doc = Document()
    doc.add_heading(title, 0)

    # Simple ToC-like index
    doc.add_heading("Contents", level=1)
    for i, sol in enumerate(solutions, start=1):
        doc.add_paragraph(f"{i}. Question ID: {sol.qid}")

    doc.add_page_break()

    # Body
    for i, sol in enumerate(solutions, start=1):
        doc.add_heading(f"{i}. Question ID: {sol.qid}", level=1)
        # Write markdown-ish content as plain paragraphs
        for para in sol.content_markdown.split("\n\n"):
            doc.add_paragraph(para)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def export_pdf(solutions: List[Solution], title: str = "UPSC Solution Book") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    flowables = []

    flowables.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
    flowables.append(Spacer(1, 12))

    # Contents
    flowables.append(Paragraph("<b>Contents</b>", styles["Heading1"]))
    for i, sol in enumerate(solutions, start=1):
        flowables.append(Paragraph(f"{i}. Question ID: {sol.qid}", styles["Normal"]))
    flowables.append(Spacer(1, 18))

    # Body
    for i, sol in enumerate(solutions, start=1):
        flowables.append(Paragraph(f"<b>{i}. Question ID: {sol.qid}</b>", styles["Heading2"]))
        # Replace newlines for basic formatting
        safe_html = sol.content_markdown.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        safe_html = safe_html.replace("\n", "<br/>")
        flowables.append(Paragraph(safe_html, styles["Normal"]))
        flowables.append(Spacer(1, 24))

    doc.build(flowables)
    return buf.getvalue()

# ---------------------------
# Sidebar Controls
# ---------------------------
st.sidebar.header("Settings")
year = st.sidebar.text_input("Year (optional)")
paper = st.sidebar.selectbox("Paper", ["Prelims", "Mains", "Other"], index=0)
process_limit = st.sidebar.number_input("Max questions to process", 1, 200, 20)
show_validator_notes = st.sidebar.checkbox("Show validator notes in output", value=True)

# ---------------------------
# Main UI
# ---------------------------
uploaded = st.file_uploader("Upload a PDF with previous year questions", type=["pdf"])

if uploaded:
    with st.spinner("Extracting text…"):
        raw_text = extract_text_from_pdf(uploaded)
    blocks = split_questions(raw_text)

    if not blocks:
        st.error("Could not detect numbered questions in this PDF. Check formatting or try another file.")
    else:
        st.success(f"Detected {len(blocks)} question blocks.")
        parsed: List[Question] = []
        for i, b in enumerate(blocks[:process_limit]):
            q = parse_mcq_block(b)
            q.year, q.paper, q.page = year, paper, i + 1
            q.topic_tags = tag_topics(q.text)
            parsed.append(q)

        st.subheader("Preview detected questions")
        for q in parsed:
            with st.expander(f"Q: {q.text[:90]}..."):
                st.markdown("**Detected Type:** " + q.qtype)
                if q.options:
                    st.markdown("**Options:**\n\n" + "\n".join(q.options))
                st.markdown("**Tags:** " + ", ".join(q.topic_tags or []))

        # Generate Solutions
        if st.button("⚙️ Generate Solutions"):
            if "OPENAI_API_KEY" not in st.secrets:
                st.error("OPENAI_API_KEY missing. Add it in Streamlit Cloud → App → Settings → Secrets.")
            else:
                results: List[Solution] = []
                with st.spinner("Generating structured solutions…"):
                    for q in parsed:
                        try:
                            sol = generate_solution(q)
                            # optionally strip validator notes
                            if not show_validator_notes:
                                sol.content_markdown = sol.content_markdown.split("\n---\nValidator notes", 1)[0]
                            results.append(sol)
                        except Exception as e:
                            st.error(f"Failed for question starting: '{q.text[:60]}…' → {e}")

                if results:
                    st.subheader("Generated Solutions")
                    for q, sol in zip(parsed, results):
                        with st.expander(f"✅ Solution: {q.text[:80]}..."):
                            st.markdown(sol.content_markdown)

                    # Exports
                    docx_bytes = export_docx(results)
                    pdf_bytes = export_pdf(results)

                    st.markdown("### 📥 Download Solution Book")
                    col1, col2 = st.columns(2)
                    with col1:
                        st.download_button(
                            label="⬇️ Download as DOCX",
                            data=docx_bytes,
                            file_name="UPSC_Solution_Book.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    with col2:
                        st.download_button(
                            label="⬇️ Download as PDF",
                            data=pdf_bytes,
                            file_name="UPSC_Solution_Book.pdf",
                            mime="application/pdf"
                        )

        # Debug download of raw text
        st.download_button(
            "⬇️ Download raw extracted text (debug)",
            data=raw_text,
            file_name="extracted_text.txt",
            mime="text/plain"
        )

else:
    st.info("Upload a UPSC/PCS PYQ PDF to begin.")


# ---------------------------
# Notes:
# - Add OPENAI_API_KEY in Streamlit Cloud Secrets:
#   OPENAI_API_KEY = "sk-..."
# - requirements.txt:
#     streamlit
#     pdfplumber
#     openai
#     python-docx
#     reportlab
# - Deploy on Streamlit Cloud → point to app.py
# ---------------------------
